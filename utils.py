import os
import logging
import re # Import regex
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT, WD_TAB_LEADER
import pypdf
# Removed tkinter imports as messagebox will be replaced by logging
# import tkinter as tk
# from tkinter import filedialog, messagebox

# Configure logging
# It's good practice to configure logging at the application entry point,
# but adding a basic config here for standalone testing/use.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
log = logging.getLogger(__name__)

# Import markers from agent script (assuming it's in the same directory)
# If not, adjust the import path accordingly
try:
    # Assuming job_application_agent defines the FMT_* constants
    import job_application_agent as agent_defs
    log.info("Successfully imported markers from job_application_agent.")
except ImportError:
    # Define markers locally as a fallback if import fails
    log.warning("Could not import markers from job_application_agent. Using fallback definitions.")
    # Simple class to mimic the structure if import fails
    class FallbackAgentDefs:
        FMT_NAME = "@@NAME@@"
        FMT_CONTACT = "@@CONTACT@@"
        FMT_HEADING = "@@HEADING@@"
        FMT_SUBHEADING_COMPANY = "@@SUBHEAD_COMP@@"
        FMT_SUBHEADING_TITLE = "@@SUBHEAD_TITLE@@"
        FMT_SUBHEADING_PROJECT = "@@SUBHEAD_PROJ@@"
        FMT_DATES = "@@DATES@@"
        FMT_BULLET = "@@BULLET@@"
        FMT_NORMAL = "@@NORMAL@@"
    agent_defs = FallbackAgentDefs()


# --- Existing Functions (parse_resume, save_text_to_file - modified to remove messagebox) ---

def parse_resume(file_path):
    """
    Parses the text content from a resume file (.pdf or .docx).
    Replaced messagebox with logging.

    Args:
        file_path (str): The path to the resume file.

    Returns:
        str: The extracted text content, or None if parsing fails or file type is unsupported.
    """
    if not file_path or not os.path.exists(file_path):
        log.error(f"File not found or path is invalid: {file_path}")
        # messagebox.showerror("Error", f"File not found or path is invalid:\n{file_path}") # Replaced
        return None

    _, file_extension = os.path.splitext(file_path)
    text_content = ""

    try:
        if file_extension.lower() == '.pdf':
            log.info(f"Parsing PDF file: {file_path}")
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                num_pages = len(reader.pages)
                log.info(f"Found {num_pages} page(s) in PDF.")
                for page_num in range(num_pages):
                    page = reader.pages[page_num]
                    extracted = page.extract_text()
                    if extracted: # Check if text was actually extracted
                         text_content += extracted + "\n"
            log.info("Successfully parsed PDF.")

        elif file_extension.lower() == '.docx':
            log.info(f"Parsing DOCX file: {file_path}")
            document = Document(file_path)
            for para in document.paragraphs:
                text_content += para.text + "\n"
            log.info("Successfully parsed DOCX.")

        else:
            log.warning(f"Unsupported file type: {file_extension}")
            # messagebox.showwarning("Unsupported File", f"Unsupported file type: {file_extension}\nPlease select a .pdf or .docx file.") # Replaced
            return None

        if not text_content.strip():
             log.warning(f"No text content extracted from file: {file_path}")
             # messagebox.showwarning("Empty File", f"Could not extract text from the file:\n{os.path.basename(file_path)}") # Replaced
             return None

        return text_content.strip()

    except Exception as e:
        log.error(f"Error parsing file {file_path}: {e}", exc_info=True)
        # messagebox.showerror("Parsing Error", f"An error occurred while parsing the file:\n{os.path.basename(file_path)}\n\nError: {e}") # Replaced
        return None

# Note: save_text_to_file still uses filedialog.asksaveasfilename.
# This function will still require a GUI environment to run unless
# filedialog is replaced or handled differently in the calling code.
# However, the messagebox calls within it are removed.
def save_text_to_file(content, initial_dir=".", title="Save File", default_extension=".txt", filetypes=[("Text files", "*.txt"), ("Word Document", "*.docx"), ("All files", "*.*")]):
    """
    Opens a save file dialog and saves the given text content to the selected file.
    Handles saving as .txt or .docx based on chosen extension.
    Replaced messagebox with logging. Requires tkinter for filedialog.

    Args:
        content (str): The text content to save.
        initial_dir (str): The initial directory for the file dialog.
        title (str): The title for the file dialog window.
        default_extension (str): The default file extension.
        filetypes (list): List of tuples defining allowed file types.

    Returns:
        str: The path where the file was saved, or None if cancelled or error.
    """
    # This part still requires a GUI environment because of filedialog
    try:
        import tkinter as tk
        from tkinter import filedialog
        root_temp = tk.Tk()
        root_temp.withdraw() # Hide the root window
        file_path = filedialog.asksaveasfilename(
            initialdir=initial_dir,
            title=title,
            defaultextension=default_extension,
            filetypes=filetypes
        )
        root_temp.destroy() # Clean up the temporary window
    except ImportError:
        log.error("Tkinter is required for the save file dialog but is not installed or available.")
        return None
    except Exception as e:
        log.error(f"Error during file dialog creation: {e}", exc_info=True)
        return None


    if not file_path:
        log.info("Save operation cancelled by user.")
        return None # User cancelled

    try:
        # Check the chosen extension to save appropriately
        _ , chosen_extension = os.path.splitext(file_path)

        if chosen_extension.lower() == '.docx':
             # Save as DOCX (plain text paragraphs)
             doc = Document()
             # Add content paragraph by paragraph to preserve some structure if needed
             for paragraph_text in content.split('\n'):
                 # Avoid adding empty paragraphs excessively, unless it's intended spacing
                 if paragraph_text.strip() or content.strip() == paragraph_text: # Add if not just whitespace, or if it's the only content
                     doc.add_paragraph(paragraph_text)
             doc.save(file_path)
             log.info(f"Content saved successfully as plain text DOCX to: {file_path}")
        else:
             # Save as plain text (or other format)
             with open(file_path, 'w', encoding='utf-8') as f:
                 f.write(content)
             log.info(f"Content saved successfully as {chosen_extension} to: {file_path}")

        # messagebox.showinfo("Success", f"File saved successfully:\n{file_path}") # Replaced
        log.info(f"File saved successfully: {file_path}")
        return file_path
    except Exception as e:
        log.error(f"Error saving file to {file_path}: {e}", exc_info=True)
        # messagebox.showerror("Save Error", f"Could not save the file.\n\nError: {e}") # Replaced
        return None

# --- UPDATED FORMATTING FUNCTION (No messagebox) ---
def format_resume_with_markers(marked_text, filename="formatted_resume.docx"):
    """
    Parses text containing specific markers (e.g., @@HEADING@@) and creates
    a formatted Word document based on predefined styles for those markers.
    Handles inline **bold** markers and formats company/title/date on the same line.
    Replaced messagebox calls with logging.

    Args:
        marked_text (str): The AI-generated resume text containing formatting markers.
        filename (str): The name for the output DOCX file.

    Returns:
        str: The path where the file was saved, or None if error.
    """
    if not marked_text or not isinstance(marked_text, str) or marked_text.strip().startswith(("(", "Error:", "(Agent Error:")):
        log.error(f"Invalid or error content passed to format_resume_with_markers. Content type: {type(marked_text)}")
        # messagebox.showerror("Formatting Error", "Cannot format resume due to invalid input content.") # Replaced
        return None

    log.info(f"Attempting to format resume and save to {filename}")
    try:
        doc = Document()

        # --- Define Standard Styles (Customize as needed) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        paragraph_format = style.paragraph_format
        paragraph_format.space_before = Pt(0)
        paragraph_format.space_after = Pt(0) # Default tight spacing
        paragraph_format.line_spacing = 1 # Slightly more than single for readability

        # --- Set Margins ---
        # Standard A4 paper width approx 8.27 inches. Letter width is 8.5 inches. Using Letter.
        page_width_inches = 8.5
        left_margin_inches = 0.5
        right_margin_inches = 0.5
        printable_width_inches = page_width_inches - left_margin_inches - right_margin_inches
        # Calculate tab stop position near the right margin
        right_tab_stop_position = Inches(printable_width_inches - 0.1) # Position slightly before the margin edge

        sections = doc.sections
        for section in sections:
            section.page_width = Inches(page_width_inches) # Ensure page width is set if needed
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(left_margin_inches)
            section.right_margin = Inches(right_margin_inches)

        # --- Process Marked Text ---
        lines = marked_text.strip().split('\n')
        i = 0
        in_education_section = False # Flag to track if processing the education section

        while i < len(lines):
            line = lines[i].strip()
            log.debug(f"Processing line {i}: '{line[:50]}...'") # Log start of line processing

            if not line:
                i += 1
                continue # Skip empty lines

            # Find the main marker and text content
            marker = None
            content = line
            # Use vars() for dynamic attribute access if agent_defs is an object/class instance
            if hasattr(agent_defs, '__dict__'):
                marker_dict = vars(agent_defs)
            elif isinstance(agent_defs, type): # Handle case where agent_defs might be the class itself
                 marker_dict = {k: v for k, v in agent_defs.__dict__.items() if not k.startswith('__')}
            else:
                 log.error("agent_defs is not an object or class with expected attributes.")
                 marker_dict = {} # Avoid error, but markers won't match

            for m_key, m_val in marker_dict.items():
                # Ensure m_val is a string before calling startswith
                if m_key.startswith("FMT_") and isinstance(m_val, str) and line.startswith(m_val):
                    marker = m_val
                    content = line[len(marker):].strip()
                    log.debug(f"  Found marker: {marker}, Content: '{content[:50]}...'")
                    break # Found the marker for this line

            # Check if we're entering or leaving the EDUCATION section based on heading content
            if marker == agent_defs.FMT_HEADING and "EDUCATION" in content.upper():
                log.debug("  Entering EDUCATION section.")
                in_education_section = True
            elif marker == agent_defs.FMT_HEADING and "EDUCATION" not in content.upper():
                 # Check if we were in education section and now leaving it
                if in_education_section:
                    log.debug("  Leaving EDUCATION section.")
                    in_education_section = False


            # --- Special handling for company+title+date on same line ---
            if marker == agent_defs.FMT_SUBHEADING_COMPANY:
                log.debug("  Applying Company/Title/Date format.")
                # Initialize paragraph with company name
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = Pt(4)  # Add some space before job entries
                pf.space_after = Pt(0)   # Tight spacing after this line

                # Add company name (bold)
                company_run = p.add_run(content)
                company_run.font.name = 'Calibri'
                company_run.font.size = Pt(11)
                company_run.bold = True

                # Look ahead for title (expected next line)
                title_content = ""
                if i + 1 < len(lines) and lines[i+1].strip().startswith(agent_defs.FMT_SUBHEADING_TITLE):
                    title_content = lines[i+1].strip()[len(agent_defs.FMT_SUBHEADING_TITLE):].strip()
                    log.debug(f"    Found title: {title_content}")
                    i += 1  # Consume title line

                # Look ahead for date (expected after title)
                date_content = ""
                if i + 1 < len(lines) and lines[i+1].strip().startswith(agent_defs.FMT_DATES):
                    date_content = lines[i+1].strip()[len(agent_defs.FMT_DATES):].strip()
                    log.debug(f"    Found date: {date_content}")
                    i += 1  # Consume date line

                # Add title if found (italic, separated by comma)
                if title_content:
                    p.add_run(", ").bold = False # Separator should not be bold/italic
                    title_run = p.add_run(title_content)
                    title_run.font.name = 'Calibri'
                    title_run.font.size = Pt(11)
                    title_run.italic = True
                    title_run.bold = False

                # Add tab and date if found (right-aligned)
                if date_content:
                    # Set up tab stop for right alignment ONLY for this paragraph
                    pf.tab_stops.clear_all() # Clear any inherited tab stops
                    pf.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                    log.debug(f"    Added RIGHT tab stop at {right_tab_stop_position}")

                    # Add tab character and date text
                    p.add_run("\t") # Add the tab character to move to the stop
                    date_run = p.add_run(date_content)
                    date_run.font.name = 'Calibri'
                    date_run.font.size = Pt(10) # Slightly smaller font for date
                    date_run.bold = False

            # --- Special handling for EDUCATION section to eliminate spacing and right-align date ---
            elif in_education_section and marker == agent_defs.FMT_NORMAL:
                log.debug("  Applying Education format (tight spacing).")
                p = doc.add_paragraph()
                pf = p.paragraph_format
                pf.space_before = Pt(0) # No space before education lines
                pf.space_after = Pt(0)  # No space after education lines

                # Process inline bold formatting for the degree/university line
                parts = re.split(r'(\*\*.*?\*\*)', content)
                for part in parts:
                    if not part: continue
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        run = p.add_run(part)
                        run.bold = False
                    run.font.name = 'Calibri'
                    run.font.size = Pt(11)

                # Look ahead for date (expected next line)
                date_content = ""
                if i + 1 < len(lines) and lines[i+1].strip().startswith(agent_defs.FMT_DATES):
                    date_content = lines[i+1].strip()[len(agent_defs.FMT_DATES):].strip()
                    log.debug(f"    Found date for education: {date_content}")
                    i += 1  # Consume date line

                # Add tab and date if found (right-aligned)
                if date_content:
                    # Set up tab stop for right alignment
                    pf.tab_stops.clear_all()
                    pf.tab_stops.add_tab_stop(right_tab_stop_position, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.SPACES)
                    log.debug(f"    Added RIGHT tab stop at {right_tab_stop_position}")

                    # Add tab character and date text
                    p.add_run("\t")
                    date_run = p.add_run(date_content)
                    date_run.font.name = 'Calibri'
                    date_run.font.size = Pt(10)
                    date_run.bold = False

            else:
                # --- Regular marker handling for other sections/lines ---
                log.debug(f"  Applying standard format for marker: {marker}")
                p = doc.add_paragraph()
                pf = p.paragraph_format

                # Apply base paragraph styles (reset spacing, apply marker-specific)
                pf.space_before = Pt(0) # Reset default
                pf.space_after = Pt(0)  # Reset default (tight)
                pf.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT # Default alignment

                base_font_size = Pt(11) # Default
                base_bold = False
                base_italic = False
                base_font_name = 'Calibri'
                base_color = None # Default color (black)

                # --- Style Application Logic based on Marker ---
                if marker == agent_defs.FMT_NAME:
                    base_font_size = Pt(16)
                    base_bold = True
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    pf.space_after = Pt(1) # Small space after name
                elif marker == agent_defs.FMT_CONTACT:
                    base_font_size = Pt(10)
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    pf.space_after = Pt(8) # Space after contact info
                elif marker == agent_defs.FMT_HEADING:
                    base_font_size = Pt(12)
                    base_bold = True
                    base_color = RGBColor(0x4F, 0x81, 0xBD) # Blue color
                    pf.space_before = Pt(10) # Space before heading
                    pf.space_after = Pt(1)   # Tight space after heading text

                    # Process the heading text first (handles potential inline bold in heading)
                    parts = re.split(r'(\*\*.*?\*\*)', content)
                    for part in parts:
                        if not part: continue
                        run = p.add_run()
                        run.font.name = base_font_name
                        run.font.size = base_font_size
                        if base_color:
                            run.font.color.rgb = base_color

                        if part.startswith('**') and part.endswith('**'):
                            run.text = part[2:-2]
                            run.bold = True # Inline bold overrides base_bold for this part
                        else:
                            run.text = part
                            run.bold = base_bold # Use heading's base bold setting

                    # Add a separate paragraph for the underline effect
                    line_p = doc.add_paragraph()
                    line_pf = line_p.paragraph_format
                    line_pf.space_before = Pt(0) # No space before the line
                    line_pf.space_after = Pt(4)  # Space after the line
                    line_run = line_p.add_run('_' * 100) # Use underscores for line
                    line_run.font.size = Pt(3)      # Very small font size
                    line_run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD) # Match heading color
                    line_run.bold = False # Underline should not be bold

                    # Skip the standard content processing below for this iteration
                    i += 1
                    continue

                elif marker == agent_defs.FMT_SUBHEADING_PROJECT:
                    base_font_size = Pt(11)
                    base_bold = True
                    pf.space_before = Pt(4) # Space before project heading
                    pf.space_after = Pt(1)  # Tight space after
                elif marker == agent_defs.FMT_DATES:
                    # This case handles dates that were NOT combined (e.g., maybe a standalone date?)
                    # Usually consumed by Company or Education logic, but handle defensively.
                    log.warning(f"  Processing standalone DATE line: {content}")
                    base_font_size = Pt(10)
                    pf.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                    pf.space_before = Pt(0)
                    pf.space_after = Pt(0)
                elif marker == agent_defs.FMT_BULLET:
                    # Apply bullet style properties directly instead of using named style 'List Bullet'
                    # This gives more control and avoids potential missing style issues.
                    pf.left_indent = Inches(0.35) # Indent text
                    pf.first_line_indent = Inches(-0.20) # Hanging indent for bullet
                    pf.space_after = Pt(2) # Space after bullet point
                    # Add bullet character if not already present
                    bullet_char = "•" # Standard bullet character
                    if not content.lstrip().startswith(('-', '*', '•')):
                         content = f"{bullet_char} {content}"
                    else:
                        # Standardize bullet char if using others
                        content = f"{bullet_char} {content.lstrip('*- ')}"

                    # Set paragraph style properties for bullet appearance
                    # (This part might need adjustment based on exact docx behavior)
                    # p.style = 'List Bullet' # Avoid using named style for robustness

                elif marker == agent_defs.FMT_NORMAL:
                    # Use default normal style settings (tight spacing)
                    pf.space_after = Pt(0)
                else:
                    # Fallback for lines without recognized markers (treat as Normal)
                    log.warning(f"Line without recognized marker, applying Normal style: '{line[:50]}...'")
                    pf.space_after = Pt(0) # Keep tight spacing

                # --- Add Text Runs (Handle Inline Bold for regular markers) ---
                parts = re.split(r'(\*\*.*?\*\*)', content)
                for part in parts:
                    if not part: continue
                    run = p.add_run()
                    run.font.name = base_font_name
                    run.font.size = base_font_size
                    run.italic = base_italic
                    if base_color:
                        run.font.color.rgb = base_color

                    # Apply bold based on marker AND inline formatting
                    if part.startswith('**') and part.endswith('**'):
                        run.text = part[2:-2]
                        run.bold = True # Inline bold always wins
                    else:
                        run.text = part
                        run.bold = base_bold # Use the marker's base bold setting

            i += 1 # Move to next line

        # --- Save Document ---
        doc.save(filename)
        log.info(f"Formatted resume saved successfully to: {filename}")
        # messagebox.showinfo("Save Successful", f"Formatted resume saved to:\n{filename}") # Replaced
        return filename

    except Exception as e:
        log.error(f"Error during formatting or saving DOCX file {filename}: {e}", exc_info=True)
        # messagebox.showerror("Save Error", f"Could not save the formatted resume file.\n\nError: {e}") # Replaced
        return None


# --- Example Usage (Requires GUI only if save_text_to_file is called) ---
if __name__ == "__main__":
    # Example usage: Does not require Tkinter root unless save_text_to_file is used.
    # root_test = tk.Tk() # No longer strictly needed for format_resume_with_markers
    # root_test.withdraw()

    print("--- Testing Formatted Save Function (using logging) ---")
    # Create sample marked text
    sample_marked_text = f"""
{agent_defs.FMT_NAME} John Doe
{agent_defs.FMT_CONTACT} john.doe@email.com | 123-456-7890 | linkedin.com/in/johndoe
{agent_defs.FMT_HEADING} SUMMARY
{agent_defs.FMT_NORMAL} A **highly motivated** individual seeking opportunity. **Very** skilled.
{agent_defs.FMT_HEADING} EXPERIENCE
{agent_defs.FMT_SUBHEADING_COMPANY} Example Inc.
{agent_defs.FMT_SUBHEADING_TITLE} Software Developer
{agent_defs.FMT_DATES} 2023 - Present
{agent_defs.FMT_BULLET} - Developed feature X using **Python** and **Java**.
{agent_defs.FMT_BULLET} Collaborated with team Y on project Z.
{agent_defs.FMT_SUBHEADING_COMPANY} Previous Job LLC
{agent_defs.FMT_SUBHEADING_TITLE} Junior Dev
{agent_defs.FMT_DATES} 2020 - 2022
{agent_defs.FMT_BULLET} * Learned **many** things.
{agent_defs.FMT_HEADING} EDUCATION
{agent_defs.FMT_NORMAL} University of Example | **BS Computer Science**
{agent_defs.FMT_DATES} 2018 - 2022
{agent_defs.FMT_NORMAL} Another University | MS Data Science
{agent_defs.FMT_DATES} 2022 - 2024
{agent_defs.FMT_HEADING} SKILLS
{agent_defs.FMT_NORMAL} Python, Java, SQL, **DOCX Formatting**, Problem Solving
    """
    # Define the output filename
    output_filename = "formatted_resume_example_log.docx"
    saved_formatted_path = format_resume_with_markers(sample_marked_text, filename=output_filename)

    if saved_formatted_path:
        print(f"Formatted file saved successfully to: {os.path.abspath(saved_formatted_path)}")
        # Optional: Try opening the file (platform dependent)
        try:
            if os.name == 'nt': # Windows
                os.startfile(os.path.abspath(saved_formatted_path))
            elif os.uname().sysname == 'Darwin': # macOS
                os.system(f'open "{os.path.abspath(saved_formatted_path)}"')
            else: # Linux
                os.system(f'xdg-open "{os.path.abspath(saved_formatted_path)}"')
        except Exception as e:
            print(f"Could not automatically open the file: {e}")
    else:
        print(f"Formatted save function failed. Check logs for errors. Attempted to save as {output_filename}")

    # root_test.destroy() # No longer needed if save_text_to_file wasn't called
